import simplepush
import requests
from tqdm import tqdm
import csv
from docx import Document
from docx.shared import Inches
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import re
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.common.exceptions import NoSuchElementException, TimeoutException
import undetected_chromedriver as uc

from seleniumbase import Driver

import pandas as pd
from time import sleep, perf_counter
import os
import dotenv
dotenv.load_dotenv()

doc_path = f"{os.environ['ROOT']}/linkedin_posts.docx"
csv_path = f"{os.environ['ROOT']}/linkedin_posts.csv"


def main():
    print("Started the program!")

    driver: WebDriver = Driver(uc=True, no_sandbox=True, headless=True)
    wait = WebDriverWait(driver, 30)
    driver.get("https://www.linkedin.com/login")

    email_input = driver.find_element(
        by=By.CSS_SELECTOR,
        value='input[id="username"]'
    )
    email_input.send_keys(os.environ["EMAIL"])

    pass_input = driver.find_element(
        by=By.CSS_SELECTOR,
        value='input[id="password"]'
    )
    pass_input.send_keys(os.environ["PASS"])

    button = driver.find_element(
        by=By.CSS_SELECTOR,
        value='button[type="submit"]'
    )
    button.click()

    sleep(3)

    driver.get("https://www.linkedin.com/in/kent-choi/recent-activity/all/")

    last_height = get_scroll_height(driver=driver)

    if not os.path.exists(doc_path):
        doc = Document()
    else:
        os.remove(doc_path)
        doc = Document()

    added_count = 0
    post_ids = set()

    while True:
        current_posts = wait.until(EC.presence_of_all_elements_located((
            By.XPATH,
            '//div[contains(@class, "scaffold-finite-scroll__content")]//div[contains(@class, "feed-shared-update-v2__description-wrapper")]'
        )))

        len_posts = len(current_posts)
        start_index = 0 if len_posts < 20 else len_posts-20
        is_end = True

        # Open the file in append mode ('a') and write the new data
        with open(csv_path, mode='a', newline='', encoding="utf-8") as file:
            writer = csv.writer(file)

            for post in current_posts[start_index:]:
                try:
                    post_text = post.find_element(
                        by=By.XPATH,
                        value=".//span[contains(@class, 'break-words')]"
                    ).text
                except:
                    post_text = ""

                post_id = post_text[0:100]
                if post_id in post_ids:
                    continue

                is_end = False
                post_ids.add(post_id)

                try:
                    images = post.find_elements(
                        by=By.XPATH,
                        value='./following-sibling::div[contains(@class, "update-components-image")]//img'
                    )
                except:
                    images = []

                image_names = []
                for i, img in enumerate(images):
                    src = img.get_attribute("src")
                    name = downloadImage(added_count + 1, i, src)
                    if name != "":
                        image_names.append(name)

                writer.writerow(
                    [f'Post number {added_count+1}', post_text, "\n".join(image_names)])

                doc.add_heading(f'Post number {added_count+1}:', level=2)
                added_count += 1
                doc.add_paragraph(post_text)
                for img_name in image_names:
                    doc.add_picture(
                        f"{os.environ['ROOT']}/images/{img_name}", width=Inches(5))

                print(post_text)
                print("----------------------------------------------------")

        if is_end:
            break

        driver.execute_script(
            'arguments[0].scrollIntoView();', current_posts[len(current_posts)-1])

        # Save the document with the new content added
        doc.save(doc_path)

        sleep(3)

        new_height = get_scroll_height(driver=driver)

        if new_height == last_height:
            break

        last_height = new_height

    doc.save(doc_path)

    simplepush.send(
        key=os.environ["SIMPLEPUSH"], title='Done scraping linkedin posts!',
        message=f'Scraped total {len(post_ids)} from linkedin account!')

    print("Done! Program closes after 1 min")


def downloadImage(index: int, image_index: int, src: str) -> str:
    response = requests.get(src)

    if response.status_code == 200:
        with open(f"{os.environ['ROOT']}/images/post_{index}_{image_index}.jpg", 'wb') as file:
            file.write(response.content)
        print("Image downloaded successfully.")
        return f"post_{index}_{image_index}.jpg"
    else:
        print("Error: Failed to download image.")

    return ""


def get_scroll_height(driver: WebDriver):
    return driver.execute_script("return document.body.scrollHeight")


def scroll_to_bottom(driver: WebDriver):
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")


if __name__ == "__main__":
    main()
